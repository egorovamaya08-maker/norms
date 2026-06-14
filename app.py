import streamlit as st
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.table import Table as DocxTable
from docx.oxml.ns import qn
import re
from collections import defaultdict, Counter
import zipfile
from lxml import etree
import unicodedata
from docx.shared import Cm, Pt, Inches

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

TABLE_CONTINUATION_RE = re.compile(
    r'^(продолжение|окончание)\s+таблицы\s+\d+', 
    re.IGNORECASE
)

# ------------------------------------------------------------
# Нормализация текста: удаление невидимок, замена пробелов
# ------------------------------------------------------------




def iter_block_items(parent):
    """
    Последовательно обходит абзацы и таблицы в документе 
    в том порядке, в котором они реально идут в файле.
    """
    from docx.document import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def normalize_text(text):
    if not text:
        return text
    result = []
    prev_digit = False
    for ch in text:
        cat = unicodedata.category(ch)
        if cat == 'Cf':          # невидимка
            prev_digit = ch.isdigit() if not result else result[-1].isdigit()
            continue
        if ch.isspace():
            result.append(' ')
        else:
            # Между цифрой и буквой нужен пробел
            if result and result[-1].isdigit() and ch.isalpha():
                result.append(' ')
            result.append(ch)
    return re.sub(r'\s+', ' ', ''.join(result)).strip()

# ------------------------------------------------------------
# Вспомогательные функции
# ------------------------------------------------------------
def get_effective_alignment(paragraph):
    if paragraph.alignment is not None:
        return paragraph.alignment
    try:
        style = paragraph.style
        if style and style.paragraph_format.alignment is not None:
            return style.paragraph_format.alignment
    except:
        pass
    return None

def is_paragraph_bold(paragraph):
    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs:
        if any(r.bold for r in runs):
            return True
    try:
        if paragraph.style and paragraph.style.font and paragraph.style.font.bold:
            return True
    except:
        pass
    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            pPr_rPr = pPr.find(qn('w:rPr'))
            if pPr_rPr is not None:
                bold_elem = pPr_rPr.find(qn('w:b'))
                if bold_elem is not None:
                    val = bold_elem.get(qn('w:val'))
                    if val != 'false' and val != '0':
                        return True
            for r in paragraph._element.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is not None:
                    bold_elem = rPr.find(qn('w:b'))
                    if bold_elem is not None:
                        val = bold_elem.get(qn('w:val'))
                        if val != 'false' and val != '0':
                            return True
    except:
        pass
    return False

def is_empty_paragraph(paragraph):
    """Надёжная проверка действительно пустого абзаца"""
    if not paragraph.text.strip():
        # Проверяем наличие любого текста в XML (включая гиперссылки и т.д.)
        full_text = "".join(node.text or "" for node in paragraph._element.iter(qn('w:t')))
        return len(full_text.strip()) == 0
    return False
# def is_empty_paragraph(paragraph):
    # Извлекаем вообще весь текст из XML-элемента абзаца, включая текст внутри w:hyperlink
#    full_text = "".join(node.text or "" for node in paragraph._element.iter(qn('w:t'))).strip()
#    return len(full_text) == 0

def has_page_number(text):
    return bool(re.search(r'[\t\s\.]{2,}\d+$', text))

def is_last_element_on_page(elem, body_elems, idx):
    """Возвращает True, если элемент находится в конце страницы (перед разрывом страницы или секции)."""
    # Смотрим следующие элементы до 5 штук
    for i in range(idx + 1, min(idx + 5, len(body_elems))):
        next_elem = body_elems[i]
        # Если нашли явный разрыв страницы
        if next_elem.tag == qn('w:p'):
            for br in next_elem.findall('.//w:br', NSMAP):
                if br.get(qn('w:type')) == 'page':
                    return True
        # Если нашли разрыв секции (новая страница)
        if next_elem.tag == qn('w:sectPr'):
            return True
        # Если встретили непустой абзац – страница не кончилась
        if next_elem.tag == qn('w:p'):
            txt = ''.join(node.text or '' for node in next_elem.iter() if node.tag == qn('w:t')).strip()
            if txt:
                return False
    return False

def is_all_caps(text):
    clean_text = re.sub(r'[\d\s\.,;:!?\-–—()«»""''«»]', '', text)
    if not clean_text:
        return False
    return clean_text == clean_text.upper()

def is_section_header(text):
    cleaned = normalize_text(text)
    if not cleaned:
        return False
    if cleaned.upper() in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ"}:
        return True
    if re.match(r'^(?:ГЛАВА|РАЗДЕЛ)\s+\d+', cleaned, re.IGNORECASE):
        return True
    # Основной критерий: номер с точкой + заглавные буквы
    if re.match(r'^\d+\.\s*[А-ЯЁ]', cleaned):
        clean_letters = re.sub(r'[\d\s\.,;:!?\-–—()«»""''«»]', '', cleaned)
        if not clean_letters:
            return False
        upper_count = sum(1 for c in clean_letters if c.isupper())
        return upper_count >= len(clean_letters) * 0.8
        # Заголовок без номера, написанный ЗАГЛАВНЫМИ БУКВАМИ
    only_letters = re.sub(r'[\d\s\.,;:!?\-–—()«»""''«»]', '', cleaned)
    if only_letters and len(only_letters) > 3 and only_letters == only_letters.upper():
        if '«' not in cleaned and '»' not in cleaned:
            return True
    return False

def normalize_title(text):
    text = re.sub(r'^(?:ГЛАВА|РАЗДЕЛ)\s+\d+[\.\s]*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^\d+(?:\.\d+)*[\s\.]+', '', text)
    return text.strip().upper()

def extract_toc_entries(doc, start_idx):
    toc_entries = []
    for i in range(start_idx):
        txt = doc.paragraphs[i].text.strip()
        if not txt:
            continue
        if has_page_number(txt):
            clean = re.sub(r'[\t\s\.]{2,}\d+$', '', txt).strip()
            if clean and len(clean) > 5:
                toc_entries.append(clean)
    return toc_entries

# ========== ВСТАВИТЬ НОВУЮ ФУНКЦИЮ ЗДЕСЬ ==========

def extract_toc_entries_clean(doc, start_idx=0):
    """
    Глубокий разбор автоматического оглавления (TOC) через XML.
    Извлекает текст строк, игнорируя служебные символы и номера страниц.
    """
    toc_entries = []
    
    # Ищем все гиперссылки в документе (Word оборачивает строки TOC в w:hyperlink для навигации)
    body_element = doc.element.body
    
    for hyperlink in body_element.xpath('.//w:hyperlink'):
        # Собираем весь текст внутри этой гиперссылки
        text_nodes = hyperlink.xpath('.//w:t')
        if not text_nodes:
            continue
            
        full_text = "".join(node.text for node in text_nodes if node.text).strip()
        
        # Если строка пустая — пропускаем
        if not full_text:
            continue
        
        # Улучшенная очистка: отделяем номер страницы от текста
        # Номер страницы может быть в конце без пробелов (например, "ВВЕДЕНИЕ5" -> "ВВЕДЕНИЕ")
        # Ищем цифры в конце строки
        clean_text = full_text
        
        # Убираем номера страниц (цифры в конце, с точками или без)
        # Шаблон: цифры и точки, которые могут быть в конце
        clean_text = re.sub(r'[\d\.]+$', '', clean_text)
        
        # Убираем отточия и пробелы в конце
        clean_text = re.sub(r'[\.\s]+$', '', clean_text)
        
        # Если после очистки текст стал слишком коротким или пустым — пропускаем
        if len(clean_text) < 2:
            continue
        
        # Дополнительная очистка: убираем номера страниц, которые могли остаться внутри
        # Например, если было "1.1 Введение 5" -> "1.1 Введение"
        clean_text = re.sub(r'\s+\d+$', '', clean_text)
        
        # Убираем множественные точки (отточия)
        clean_text = re.sub(r'\.{2,}', '', clean_text)
        
        # Убираем табы и лишние пробелы
        clean_text = re.sub(r'[\t\s]+', ' ', clean_text).strip()
        
        if clean_text and len(clean_text) > 2 and not clean_text.isdigit():
            if clean_text not in toc_entries:
                toc_entries.append(clean_text)
    
    # Если XML-метод через гиперссылки не сработал, используем резервный метод
    if not toc_entries:
        search_limit = min(50, len(doc.paragraphs))
        for i in range(search_limit):
            txt = doc.paragraphs[i].text.strip()
            if not txt:
                continue
            # Убираем номера страниц из текста
            clean = re.sub(r'\s+\d+$', '', txt)
            clean = re.sub(r'[\t\s\.]{2,}\d+$', '', clean)
            clean = re.sub(r'\.{2,}', '', clean)
            clean = clean.strip()
            
            # Проверяем, похоже ли на строку оглавления
            if (clean and len(clean) > 3 and 
                (re.match(r'^\d+(\.\d+)*\s+[А-Яа-я]', clean) or 
                 clean.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"])):
                if clean not in toc_entries:
                    toc_entries.append(clean)
    
    return toc_entries  


def is_dash_char(ch):
    code = ord(ch)
    if code in [0x2D, 0x2010, 0x2011, 0x2012, 0x2013, 0x2014, 0x2015]:
        return True
    if code in [0xF02D]:
        return True
    return False

def is_bullet_char(ch):
    code = ord(ch)
    return code in [0x2022, 0x2023, 0x25E6, 0x25CF, 0x25CB, 0x26AB, 0x2B24]

def is_arrow_char(ch):
    code = ord(ch)
    return (0x2190 <= code <= 0x21FF or
            0x2794 <= code <= 0x27BF or
            0x2B00 <= code <= 0x2BFF)

def get_list_marker_info(paragraph, doc):
    text = paragraph.text.strip()
    if not text:
        return False, "", True

    first_char = text[0] if text else ""

    if is_dash_char(first_char):
        return True, "тире", True

    if re.match(r'^\d+\)\s', text):
        return True, "нумерованный", True
    if re.match(r'^\d+\.\s', text):
        return True, "нумерованный", True

    if re.match(r'^[а-яё]\)\s', text, re.IGNORECASE):
        return True, "буквенный", True
    if re.match(r'^[a-z]\)\s', text, re.IGNORECASE):
        return True, "буквенный", True

    if is_bullet_char(first_char):
        return True, "круглый маркер (•)", False
    if is_arrow_char(first_char):
        return True, f"недопустимый маркер (стрелка {first_char})", False

    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_elem = numPr.find(qn('w:numId'))
                if numId_elem is not None:
                    numId = numId_elem.get(qn('w:val'))
                    numbering_part = doc.part.numbering_part
                    if numbering_part is not None:
                        numbering_xml = numbering_part._element
                        abstractNumId = None
                        for num in numbering_xml.findall(qn('w:num')):
                            if num.get(qn('w:numId')) == numId:
                                aid_elem = num.find(qn('w:abstractNumId'))
                                if aid_elem is not None:
                                    abstractNumId = aid_elem.get(qn('w:val'))
                                break
                        if abstractNumId:
                            for abstractNum in numbering_xml.findall(qn('w:abstractNum')):
                                if abstractNum.get(qn('w:abstractNumId')) == abstractNumId:
                                    for lvl in abstractNum.findall(qn('w:lvl')):
                                        numFmt = lvl.find(qn('w:numFmt'))
                                        lvlText_elem = lvl.find(qn('w:lvlText'))
                                        if numFmt is not None:
                                            fmt = numFmt.get(qn('w:val'))
                                            if fmt == 'bullet':
                                                if lvlText_elem is not None:
                                                    txt_val = lvlText_elem.get(qn('w:val'))
                                                    if txt_val:
                                                        clean = re.sub(r'%\d+', '', txt_val).strip()
                                                        if clean and is_dash_char(clean[0]):
                                                            return True, "тире", True
                                                        else:
                                                            return True, "круглый маркер (•)", False
                                                else:
                                                    return True, "круглый маркер (•)", False
                                            elif fmt == 'decimal':
                                                return True, "нумерованный (цифры)", True
                                            elif fmt in ['lowerLetter', 'upperLetter']:
                                                return True, "буквенный", True
                                            else:
                                                return True, f"формат '{fmt}'", True
    except:
        pass

    left_indent = get_effective_left_indent(paragraph)
    first_line = get_effective_first_line_indent(paragraph)
    if left_indent > 0.5 and first_line < 0:
        if not re.match(r'[А-Яа-яёЁA-Za-z0-9]', first_char) and first_char != ' ':
            return True, f"недопустимый маркер ({first_char})", False

    return False, "", True

def get_effective_first_line_indent(paragraph):
    pf = paragraph.paragraph_format
    # 1. Если у абзаца ЕСТЬ собственный (локальный) отступ, возвращаем его
    if pf.first_line_indent is not None:
        return pf.first_line_indent.cm

    # 2. Если локального отступа нет, берем отступ из стиля абзаца или его родительских стилей
    try:
        current_style = paragraph.style
        while current_style is not None:
            if current_style.paragraph_format.first_line_indent is not None:
                return current_style.paragraph_format.first_line_indent.cm
            current_style = current_style.base_style
    except:
        pass

    # 3. Глубокий поиск в XML (на случай, если свойства заданы через твипы напрямую в pPr)
    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                first_line = ind.get(qn('w:firstLine'))
                if first_line is not None:
                    return int(first_line) / 567  # Перевод твипов в см
    except:
        pass

    return 0.0

def get_effective_left_indent(paragraph):
    pf = paragraph.paragraph_format
    # 1. Локальный левый отступ
    if pf.left_indent is not None:
        return pf.left_indent.cm
        
    # 2. Левый отступ из стилей
    try:
        current_style = paragraph.style
        while current_style is not None:
            if current_style.paragraph_format.left_indent is not None:
                return current_style.paragraph_format.left_indent.cm
            current_style = current_style.base_style
    except:
        pass
        
    # 3. Левый отступ из XML
    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                left = ind.get(qn('w:left'))
                if left is not None:
                    return int(left) / 567
    except:
        pass
        
    return 0.0

def is_table_continuation(text):
    return bool(re.match(r'^(?:Продолжение|Окончание)\s+таблицы?\s*\d', text))

def get_table_depth(table):
    depth = 0
    element = table._element
    parent = element.getparent()
    while parent is not None:
        if parent.tag == qn('w:tbl'):
            depth += 1
        parent = parent.getparent()
    return depth

def is_formula_or_equation(text):
    if re.search(r'[=≠≤≥±×÷∫∑∏√∞∂∇∈∉⊂⊃∪∩]', text):
        return True
    if re.search(r'[α-ωΑ-Ω]', text):
        return True
    if re.search(r'[₀-₉ₐ-ₜₓᵦ-ᵧ⁰-⁹ⁱ⁻⁺ⁿ]', text):
        return True
    return False

def is_formula_where_line(text):
    return bool(re.match(r'^[Гг]де\s*[:\s]?\s*[А-Яа-яA-Za-z\-–—]', text))

def check_formula_explanation(text, paragraph, prev_was_formula, prev_para_empty):
    errors = []
    short_text = text[:60] + "…" if len(text) > 60 else text
    key = f"Пояснение к формуле «{short_text}»"
    if text.startswith("Где"):
        errors.append(f"{key} – «где» должно быть с маленькой буквы")
    if re.match(r'^[Гг]де\s*:', text):
        errors.append(f"{key} – уберите двоеточие после «где»")
    first_line = get_effective_first_line_indent(paragraph)
    if abs(first_line - 1.0) > 0.2:
        errors.append(f"{key} – установите абзацный отступ 1,0 см (сейчас {first_line:.1f} см)")
    if prev_para_empty and prev_was_formula:
        errors.append(f"{key} – уберите пустую строку перед пояснением (должно идти сразу после формулы)")
    explanation_text = re.sub(r'^[Гг]де\s*:?\s*', '', text).strip()
    lines = [l.strip() for l in explanation_text.split('\t') if l.strip()]
    if len(lines) > 1:
        for i, line in enumerate(lines):
            if i < len(lines) - 1:
                if not line.endswith(';'):
                    var_match = re.match(r'^([^\s–—]+)\s*[–—]\s*(.+)$', line)
                    var_name = var_match.group(1) if var_match else line[:20]
                    errors.append(f"{key} – после «{var_name}» должна быть точка с запятой (;)")
            else:
                if not line.endswith('.'):
                    var_match = re.match(r'^([^\s–—]+)\s*[–—]\s*(.+)$', line)
                    var_name = var_match.group(1) if var_match else line[:20]
                    errors.append(f"{key} – после «{var_name}» должна быть точка (.)")
    else:
        parts = re.findall(r'([^\s–—]+)\s*[–—]\s*([^;.]+)([;.]?)', explanation_text)
        if len(parts) > 1:
            for i, (var, desc, punct) in enumerate(parts):
                if i < len(parts) - 1:
                    if punct != ';':
                        errors.append(f"{key} – после «{var}» должна быть точка с запятой (;)")
                else:
                    if punct != '.':
                        errors.append(f"{key} – после «{var}» должна быть точка (.)")
        elif len(parts) == 1:
            var, desc, punct = parts[0]
            if punct != '.':
                errors.append(f"{key} – после «{var}» должна быть точка (.)")
    return errors

def check_page_numbering(file, intro_start_idx):
    issues = []
    try:
        with zipfile.ZipFile(file, 'r') as z:
            doc_xml = etree.fromstring(z.read('word/document.xml'))
            nsmap = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }
            body = doc_xml.find('.//w:body', nsmap)
            body_elements = list(body)
            if intro_start_idx >= len(body_elements):
                return []
            target_sectPr = None
            for i in range(intro_start_idx, len(body_elements)):
                if body_elements[i].tag == qn('w:sectPr'):
                    target_sectPr = body_elements[i]
                    break
            if target_sectPr is None:
                target_sectPr = body.find('.//w:sectPr', nsmap)
            if target_sectPr is None:
                return []
            footer_refs = target_sectPr.findall('.//w:footerReference', nsmap)
            if not footer_refs:
                return []
            try:
                rels_xml = etree.fromstring(z.read('word/_rels/document.xml.rels'))
                rels_map = {}
                for rel in rels_xml:
                    rel_id = rel.get('Id')
                    target = rel.get('Target')
                    rel_type = rel.get('Type')
                    if 'footer' in rel_type.lower():
                        rels_map[rel_id] = target
            except:
                return []
            for footer_ref in footer_refs:
                ref_id = footer_ref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                file_name = rels_map.get(ref_id)
                if not file_name:
                    continue
                full_path = 'word/' + file_name
                try:
                    content = z.read(full_path)
                    tree = etree.fromstring(content)
                    paragraphs = []
                    for para in tree.findall('.//w:p', nsmap):
                        texts = []
                        for t in para.findall('.//w:t', nsmap):
                            if t.text:
                                texts.append(t.text)
                        full_text = ''.join(texts)
                        has_page = False
                        for fld in para.findall('.//w:fldChar', nsmap):
                            fld_type = fld.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType')
                            if fld_type == 'begin':
                                has_page = True
                        for instr in para.findall('.//w:instrText', nsmap):
                            if instr.text and 'PAGE' in instr.text:
                                has_page = True
                        jc = para.find('.//w:jc', nsmap)
                        align = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if jc is not None else None
                        font_size = None
                        if has_page:
                            for r in para.findall('.//w:r', nsmap):
                                rPr = r.find('.//w:rPr', nsmap)
                                if rPr is not None:
                                    sz = rPr.find('.//w:sz', nsmap)
                                    if sz is None:
                                        sz = rPr.find('.//w:szCs', nsmap)
                                    if sz is not None:
                                        sz_val = sz.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                        if sz_val:
                                            font_size = int(sz_val) / 2
                                            break
                        paragraphs.append({
                            'text': full_text,
                            'has_page': has_page,
                            'alignment': align,
                            'font_size': font_size,
                            'is_empty': not bool(full_text)
                        })
                    has_page_in_footer = any(p['has_page'] for p in paragraphs)
                    if not has_page_in_footer:
                        static_numbers = [p for p in paragraphs if p['text'].isdigit()]
                        if static_numbers:
                            issues.append("Нумерация страниц – используйте автонумерацию (поле PAGE) вместо статичного номера")
                        else:
                            issues.append("Нумерация страниц – отсутствует нумерация в колонтитуле")
                        return issues
                    page_para_errors = []
                    for j, p in enumerate(paragraphs):
                        if p['has_page']:
                            if p['font_size'] is not None and abs(p['font_size'] - 14) > 0.5:
                                page_para_errors.append(f"размер шрифта {p['font_size']} pt (должен быть 14)")
                            if p['alignment'] != 'center':
                                page_para_errors.append("выравнивание должно быть по центру")
                            if j + 1 < len(paragraphs) and paragraphs[j + 1]['is_empty']:
                                page_para_errors.append("уберите пустую строку после номера")
                        elif p['is_empty'] and j > 0 and paragraphs[j-1]['has_page']:
                            page_para_errors.append("уберите пустую строку после номера")
                    if page_para_errors:
                        issues.append(f"Нумерация страниц – {', '.join(page_para_errors)}")
                    return issues
                except:
                    continue
    except Exception:
        pass
    return issues

def get_category_order(issue):
    if issue.startswith("Таблицы –"):
        return (4, 0)
    if issue.startswith("Рисунки –"):
        return (3, 0)
    if issue.startswith("«") and "» –" in issue:
        return (1, 1)
    if issue.startswith("Подраздел «"):
        return (2, 1)
    if issue.startswith("Рисунок "):
        return (3, 1)
    if issue.startswith("Таблица "):
        return (4, 1)
    if issue.startswith("Пояснение к формуле «"):
        return (5, 1)
    if issue.startswith("Список, начиная с «"):
        return (6, 1)
    if issue.startswith("Нумерация страниц"):
        return (7, 1)
    if issue.startswith("Поля страниц"):
        return (8, 1)
    if issue.startswith("Основной текст, начиная со строки «"):
        return (9, 1)
    if issue.startswith("Список источников"):
        return (10, 1)
    return (11, 1)

def group_issues(issues_list):
    auto_issues = []
    manual_issues = []
    manual_checks = []
    manual_section = False

    for issue in issues_list:
        if issue.startswith("📋 Для проверки человеком"):
            manual_section = True
            continue
        if manual_section:
            manual_checks.append(issue)
        else:
            auto_issues.append(issue)

    grouped = defaultdict(list)
    standalone = []

    for issue in auto_issues:
        match = re.match(
            r'^(?:«([^»]+)»|(Рисунок\s+[\d.]+)|(Таблица\s+[\d.]+)|'
            r'(Подраздел\s+«([^»]+)»)|(Пояснение к формуле «([^»]+)»)|'
            r'(Список начиная с «([^»]+)»)|(Нумерация страниц))\s*[–-]\s*(.+)$',
            issue
        )
        if match:
            key = None
            if match.group(1):
                key = match.group(1)[:80]
            elif match.group(2):
                key = match.group(2)
            elif match.group(3):
                key = match.group(3)
            elif match.group(5):
                key = f"Подраздел «{match.group(5)[:50]}»"
            elif match.group(6):
                key = f"Пояснение к формуле «{match.group(7)[:60]}»"
            elif match.group(8):
                key = f"Список начиная с «{match.group(9)[:50]}»"
            elif match.group(10):
                key = "Нумерация страниц"
            if key:
                message = match.group(11)
                grouped[key].append(message)
            else:
                standalone.append(issue)
        else:
            standalone.append(issue)

    figure_keys = [k for k in grouped if re.match(r'^Рисунок\s+\d', k)]
    if len(figure_keys) > 3:
        all_fig_messages = []
        for k in figure_keys:
            all_fig_messages.extend(grouped[k])
        msg_counts = Counter(all_fig_messages)
        unique_msgs = list(msg_counts.keys())
        first_seen = {}
        for msg in all_fig_messages:
            if msg not in first_seen:
                first_seen[msg] = len(first_seen)
        unique_msgs.sort(key=lambda m: (-msg_counts[m], first_seen[m]))
        replacements = {
            "удалите точку в конце": "в названии не должно быть точки в конце",
            "уберите точку в конце": "в названии не должно быть точки в конце"
        }
        transformed = [replacements.get(m, m) for m in unique_msgs]
        combined_fig = "Рисунки – " + ", ".join(transformed)
        for k in figure_keys:
            del grouped[k]
        standalone.insert(0, combined_fig)

    caption_msg_pattern = re.compile(r'^Исправьте название на «Таблица [\d.]+ – Название»$')
    caption_msgs_keys = []
    for key in list(grouped.keys()):
        if key.startswith("Таблица ") and re.match(r'^Таблица\s+\d', key):
            msgs = grouped[key]
            new_msgs = []
            for m in msgs:
                if caption_msg_pattern.match(m):
                    caption_msgs_keys.append(key)
                else:
                    new_msgs.append(m)
            if new_msgs:
                grouped[key] = new_msgs
            else:
                del grouped[key]

    if len(caption_msgs_keys) >= 2:
        for k in caption_msgs_keys:
            if k in grouped and not grouped[k]:
                del grouped[k]
        standalone.insert(0, "Таблицы – исправьте название на «Таблица № – Название»")
    else:
        for k in caption_msgs_keys:
            grouped[k].append("Исправьте название на «Таблица " + k.split()[-1] + " – Название»")

# ==================== ЗАМЕНИТЬ ЭТОТ БЛОК ====================
    # --- НАЧАЛО ИСПРАВЛЕНИЯ ДЛЯ ПРОВЕРКИ РАСПОЛОЖЕНИЯ НАЗВАНИЯ ТАБЛИЦЫ ---
    # 1. Проверяем сообщения, которые попали в standalone (без конкретных номеров)
    before_msgs = [issue for issue in standalone if "название должно быть перед таблицей" in issue]
    if len(before_msgs) >= 2:
        for msg in before_msgs:
            if msg in standalone:
                standalone.remove(msg)
        standalone.insert(0, "Таблицы – название должно быть перед таблицей")

    # 2. Проверяем сообщения, которые привязались к конкретным номерам таблиц (например, "Таблица 2")
    tables_with_before_error = []
    for k, v in list(grouped.items()):
        if k.startswith("Таблица ") and any("название должно быть перед таблицей" in m for m in v):
            tables_with_before_error.append(k)
            
    # Если таких таблиц много (>= 2), выносим в общее правило "Таблицы – ..."
    if len(tables_with_before_error) >= 2:
        for k in tables_with_before_error:
            # Удаляем только конкретное сообщение из списка ошибок этой таблицы
            grouped[k] = [m for m in grouped[k] if "название должно быть перед таблицей" not in m]
            # Если у таблицы больше не осталось ошибок, удаляем её из группы совсем
            if not grouped[k]:
                del grouped[k]
        if "Таблицы – название должно быть перед таблицей" not in standalone:
            standalone.insert(0, "Таблицы – название должно быть перед таблицей")

    table_common_msgs = [issue for issue in standalone if issue.startswith("Таблицы –")]
    if len(table_common_msgs) > 1:
        parts = [issue[len("Таблицы –"):] for issue in table_common_msgs]
        combined = "Таблицы – " + ", ".join(parts)
        for issue in table_common_msgs:
            standalone.remove(issue)
        standalone.append(combined)

    result = []
    for issue in standalone:
        result.append(issue)
    for key, messages in grouped.items():
        unique_msgs = list(dict.fromkeys(messages))
        if len(unique_msgs) == 1:
            if key.startswith("Подраздел «") or key.startswith("Пояснение к формуле «") or key.startswith("Список начиная с «"):
                result.append(f"{key} – {unique_msgs[0]}")
            elif key.startswith("Рисунок ") or key.startswith("Таблица ") or key == "Нумерация страниц":
                result.append(f"{key} – {unique_msgs[0]}")
            else:
                result.append(f"«{key}» – {unique_msgs[0]}")
        else:
            combined = "; ".join(unique_msgs)
            if key.startswith("Подраздел «") or key.startswith("Пояснение к формуле «") or key.startswith("Список начиная с «"):
                result.append(f"{key} – {combined}")
            elif key.startswith("Рисунок ") or key.startswith("Таблица ") or key == "Нумерация страниц":
                result.append(f"{key} – {combined}")
            else:
                result.append(f"«{key}» – {combined}")

    result.sort(key=lambda x: (get_category_order(x), x))

    cont_pattern = re.compile(r'^Таблица\s+(\d+(?:\.\d+)?)\s+–\s+проверьте\s+наличие\s+«Продолжение таблицы \1»\s+/\s+«Окончание таблицы \1» при переносе на следующую страницу$')
    cont_msgs = [m for m in manual_issues if cont_pattern.match(m)]
    if len(cont_msgs) >= 2:
        for m in cont_msgs:
            manual_issues.remove(m)
        manual_issues.append("Таблицы – проверьте наличие «Продолжение таблицы» / «Окончание таблицы» при переносе на следующую страницу")

    if manual_issues:
        result.append("\n📋 Для проверки человеком:")
        result.extend(manual_issues)

    return result

def get_font_size_pt(paragraph):
    sizes = set()
    for run in paragraph.runs:
        if run.font.size:
            sizes.add(run.font.size.pt)
    return sizes

def check_empty_line_before_after(doc, idx, start_idx, label):
    errors = []
    body_elems = list(doc.element.body)
   
    if idx > start_idx:
        prev_elem = body_elems[idx - 1]
        is_empty_prev = (
            prev_elem.tag == qn('w:p') and
            not prev_elem.xpath('string(.)').strip()
        )
        if not is_empty_prev:
            errors.append(f"{label} – добавьте пустую строку перед подписью")
    if idx + 1 < len(body_elems):
        # Если подпись – последнее на странице, пустая строка не нужна
        if is_last_element_on_page(body_elems[idx], body_elems, idx):
            return errors   # ничего не добавляем

        next_elem = body_elems[idx + 1]
        is_empty_next = (
            next_elem.tag == qn('w:p') and
            not next_elem.xpath('string(.)').strip()
        )
        if not is_empty_next:
            errors.append(f"{label} – добавьте пустую строку после подписи")
    return errors

def is_subsection_header(text):
    if re.match(r'^\d+\.\d+(\.\d+)?\s+', text.strip()):
        return True
    return False
    
def is_on_new_page(doc, body_idx, start_body_pos=0, min_empty_paragraphs=10):
    body_elems = list(doc.element.body)
    if body_idx == start_body_pos:
        return True
    
    # 1. Явные признаки разрыва страницы (XML)
    for i in range(body_idx, start_body_pos - 1, -1):
        if i < 0:
            break
        elem = body_elems[i]
        if elem.tag == qn('w:sectPr'):
            if i == len(body_elems) - 1:
                continue
            type_el = elem.find(qn('w:type'))
            val = type_el.get(qn('w:val')) if type_el is not None else None
            if val != 'continuous':
                return True
        if elem.tag == qn('w:p'):
            for br in elem.findall('.//w:br', NSMAP):
                if br.get(qn('w:type')) == 'page':
                    return True
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                if pPr.find(qn('w:pageBreakBefore')) is not None:
                    return True
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    type_el = sectPr.find(qn('w:type'))
                    val = type_el.get(qn('w:val')) if type_el is not None else None
                    if val != 'continuous':
                        return True
    
    # 2. Косвенные признаки: если перед элементом мало текста с начала предыдущего крупного заголовка
    current_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p._element == body_elems[body_idx]:
            current_idx = i
            break
    
    if current_idx is not None and current_idx > 0:
        prev_section_idx = None
        for i in range(current_idx - 1, -1, -1):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if text and (is_section_header(text) or is_subsection_header(text)):
                prev_section_idx = i
                break
        
        if prev_section_idx is not None:
            non_empty_count = 0
            for i in range(prev_section_idx + 1, current_idx):
                if doc.paragraphs[i].text.strip():
                    non_empty_count += 1
            if non_empty_count < 15:
                return True
    
    # 3. Много пустых строк перед элементом
    blank_count = 0
    for i in range(body_idx - 1, start_body_pos - 1, -1):
        elem = body_elems[i]
        if elem.tag == qn('w:p'):
            if has_content(elem):
                if blank_count >= min_empty_paragraphs:
                    return True
                break
            else:
                blank_count += 1
                continue
        if elem.tag == qn('w:tbl'):
            if blank_count >= min_empty_paragraphs:
                return True
            break
    
    return False
    
    # 1. Явные признаки разрыва страницы (XML)
    for i in range(body_idx, start_body_pos - 1, -1):
        if i < 0:
            break
        elem = body_elems[i]
        if elem.tag == qn('w:sectPr'):
            if i == len(body_elems) - 1:
                continue
            type_el = elem.find(qn('w:type'))
            val = type_el.get(qn('w:val')) if type_el is not None else None
            if val != 'continuous':
                return True
        if elem.tag == qn('w:p'):
            for br in elem.findall('.//w:br', NSMAP):
                if br.get(qn('w:type')) == 'page':
                    return True
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                if pPr.find(qn('w:pageBreakBefore')) is not None:
                    return True
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    type_el = sectPr.find(qn('w:type'))
                    val = type_el.get(qn('w:val')) if type_el is not None else None
                    if val != 'continuous':
                        return True
    
    # 2. Косвенные признаки: если перед элементом мало текста с начала предыдущего крупного заголовка
    current_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p._element == body_elems[body_idx]:
            current_idx = i
            break
    
    if current_idx is not None and current_idx > 0:
        prev_section_idx = None
        for i in range(current_idx - 1, -1, -1):
            p = doc.paragraphs[i]
            text = p.text.strip()
            # ИСПРАВЛЕНИЕ: Ловим и разделы, и подразделы, чтобы точнее понимать контекст страницы
            if text and (is_section_header(text) or is_subsection_header(text)):
                prev_section_idx = i
                break
        
        if prev_section_idx is not None:
            non_empty_count = 0
            for i in range(prev_section_idx + 1, current_idx):
                if doc.paragraphs[i].text.strip():
                    non_empty_count += 1
            if non_empty_count < 15:
                return True
    
    # 3. Много пустых строк перед элементом
    blank_count = 0
    for i in range(body_idx - 1, start_body_pos - 1, -1):
        elem = body_elems[i]
        if elem.tag == qn('w:p'):
            if has_content(elem):
                if blank_count >= min_empty_paragraphs:
                    return True
                break
            else:
                blank_count += 1
                continue
        if elem.tag == qn('w:tbl'):
            if blank_count >= min_empty_paragraphs:
                return True
            break
    
    return False
    
    # 1. Явные признаки разрыва страницы
    for i in range(body_idx, start_body_pos - 1, -1):
        if i < 0:
            break
        elem = body_elems[i]
        if elem.tag == qn('w:sectPr'):
            if i == len(body_elems) - 1:
                continue
            type_el = elem.find(qn('w:type'))
            val = type_el.get(qn('w:val')) if type_el is not None else None
            if val != 'continuous':
                return True
        if elem.tag == qn('w:p'):
            for br in elem.findall('.//w:br', NSMAP):
                if br.get(qn('w:type')) == 'page':
                    return True
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                if pPr.find(qn('w:pageBreakBefore')) is not None:
                    return True
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    type_el = sectPr.find(qn('w:type'))
                    val = type_el.get(qn('w:val')) if type_el is not None else None
                    if val != 'continuous':
                        return True
    
    # 2. Косвенные признаки: если элемент является заголовком раздела
    #    и перед ним мало непустого текста (вероятно, новая страница)
    current_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p._element == body_elems[body_idx]:
            current_idx = i
            break
    
    if current_idx is not None and current_idx > 0:
        prev_section_idx = None
        for i in range(current_idx - 1, -1, -1):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if text and is_section_header(text):
                prev_section_idx = i
                break
        
        if prev_section_idx is not None:
            non_empty_count = 0
            for i in range(prev_section_idx + 1, current_idx):
                if doc.paragraphs[i].text.strip():
                    non_empty_count += 1
            if non_empty_count < 15:
                return True
    
    # 3. Много пустых строк перед элементом
    blank_count = 0
    for i in range(body_idx - 1, start_body_pos - 1, -1):
        elem = body_elems[i]
        if elem.tag == qn('w:p'):
            if has_content(elem):
                if blank_count >= min_empty_paragraphs:
                    return True
                break
            else:
                blank_count += 1
                continue
        if elem.tag == qn('w:tbl'):
            if blank_count >= min_empty_paragraphs:
                return True
            break
    
    return False

def has_content(elem):
    texts = [node.text or '' for node in elem.iter() if node.tag == qn('w:t')]
    if any(t.strip() for t in texts):
        return True
    if elem.find('.//w:drawing', NSMAP) is not None:
        return True
    return False

def has_proper_spacing_after_header(doc, header_idx: int) -> bool:
    """
    Надёжная комплексная проверка отступа после заголовка 1 уровня.
    Сочетает lookahead-сканирование последующих абзацев и многоуровневый
    анализ свойств абзацев (Python API, Стили, OpenXML).
    """
    if header_idx + 1 >= len(doc.paragraphs):
        return False

    curr_p = doc.paragraphs[header_idx]

    # =========================================================================
    # БЛОК 1. Проверка интервала ПОСЛЕ самого заголовка (space_after)
    # =========================================================================
    # 1.1. Проверка через высокоуровневое API python-docx (явный интервал или стиль)
    try:
        if curr_p.paragraph_format.space_after and curr_p.paragraph_format.space_after.pt >= 8:
            return True
        if (curr_p.style and curr_p.style.paragraph_format.space_after and 
            curr_p.style.paragraph_format.space_after.pt >= 8):
            return True
    except:
        pass

    # 1.2. Проверка низкоуровневого OpenXML (w:spaceAfter локального переопределения)
    try:
        pPr = curr_p._element.find(qn('w:pPr'))
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                after = spacing.get(qn('w:after'))
                # 160 dxa = 8 pt (1 pt = 20 dxa)
                if after and int(after) >= 160:
                    return True
    except:
        pass


    # =========================================================================
    # БЛОК 2. Lookahead-сканирование последующих элементов (окно в 4 абзаца)
    # =========================================================================
    # Ищем либо физический пустой абзац, либо текст с отступом ПЕРЕД ним (space_before)
    for i in range(header_idx + 1, min(header_idx + 5, len(doc.paragraphs))):
        next_p = doc.paragraphs[i]
        
        # 2.1. Физический пустой абзац (перевод строки) — легитимный разделитель
        if is_empty_paragraph(next_p):
            return True

        # 2.2. Высокуровневая проверка интервала ПЕРЕД следующим абзацем (API + стили)
        try:
            pf = next_p.paragraph_format
            if pf.space_before and pf.space_before.pt >= 8:
                return True
            if (next_p.style and next_p.style.paragraph_format.space_before and 
                next_p.style.paragraph_format.space_before.pt >= 8):
                return True
        except:
            pass

        # 2.3. Низкоуровневая OpenXML-проверка интервала ПЕРЕД следующим абзацем
        try:
            pPr = next_p._element.find(qn('w:pPr'))
            if pPr is not None:
                spacing = pPr.find(qn('w:spacing'))
                if spacing is not None:
                    before = spacing.get(qn('w:before'))
                    if before and int(before) >= 160:   # 160 dxa = 8 pt
                        return True
        except:
            pass

    return False

def extract_toc_from_xml(doc):
    """Извлекает параграфы содержания из поля TOC (автособираемое оглавление)"""
    toc_paragraphs = []
    body_elems = list(doc.element.body)
    in_toc_field = False
    
    for elem in body_elems:
        # Ищем начало поля TOC
        if elem.tag == qn('w:p'):
            fldChar = elem.find('.//w:fldChar', NSMAP)
            if fldChar is not None and fldChar.get(qn('w:fldCharType')) == 'begin':
                # Проверяем, что это действительно TOC
                parent = elem.getparent()
                idx = body_elems.index(elem)
                if idx + 1 < len(body_elems):
                    next_elem = body_elems[idx + 1]
                    instr = next_elem.find('.//w:instrText', NSMAP)
                    if instr is not None and instr.text and 'TOC' in instr.text.upper():
                        in_toc_field = True
                        continue
        
        if in_toc_field:
            # Если встретили конец поля
            if elem.tag == qn('w:p'):
                end_fldChar = elem.find('.//w:fldChar', NSMAP)
                if end_fldChar is not None and end_fldChar.get(qn('w:fldCharType')) == 'end':
                    break
            
            # Если это параграф с текстом, извлекаем его
            if elem.tag == qn('w:p'):
                # Ищем соответствующий параграф в doc.paragraphs
                for p in doc.paragraphs:
                    if p._element is elem:
                        toc_paragraphs.append(p)
                        break
    
    return toc_paragraphs

def check_toc(doc, start_idx):
    """Проверка содержания и подготовка границ документа"""
    errors = []
    body_elems = list(doc.element.body)

    # === 1. Поиск заголовка «СОДЕРЖАНИЕ» ===
    toc_header_idx = None
    toc_header_para = None

    # Сначала ищем в абзацах
    for i, p in enumerate(doc.paragraphs):
        raw = "".join(node.text or "" for node in p._element.iter())
        clean = re.sub(r'[\s\u00a0\u200b\ufeff]+', '', raw.upper())
        if "СОДЕРЖАНИЕ" in clean or "ОГЛАВЛЕНИЕ" in clean:
            toc_header_idx = i
            toc_header_para = p
            break

    # Если не нашли в абзацах — ищем в таблицах
    if toc_header_para is None:
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        raw = "".join(node.text or "" for node in p._element.iter())
                        clean = re.sub(r'[\s\u00a0\u200b\ufeff]+', '', raw.upper())
                        if "СОДЕРЖАНИЕ" in clean or "ОГЛАВЛЕНИЕ" in clean:
                            toc_header_para = p
                            toc_header_idx = 0 
                            break
                    if toc_header_para: break
                if toc_header_para: break
            if toc_header_para: break

    if toc_header_para is None:
        errors.append("Содержание – отсутствует заголовок «СОДЕРЖАНИЕ»")
        return errors

    # Проверка текста заголовка
    header_text = toc_header_para.text.strip()
    if header_text.upper() == "ОГЛАВЛЕНИЕ":
        errors.append("Содержание – замените «ОГЛАВЛЕНИЕ» на «СОДЕРЖАНИЕ»")
    
    # Выравнивание заголовка «СОДЕРЖАНИЕ» по центру
    if get_effective_alignment(toc_header_para) != WD_ALIGN_PARAGRAPH.CENTER:
        errors.append("Содержание – выровняйте заголовок «СОДЕРЖАНИЕ» по центру")
    
    # Размер шрифта заголовка
    font_sizes = get_font_size_pt(toc_header_para)
    if font_sizes and any(abs(s - 14) > 0.5 for s in font_sizes):
        errors.append("Содержание – установите размер шрифта 14 пт для заголовка")
    
    # Отступы заголовка
    if abs(get_effective_first_line_indent(toc_header_para)) > 0.1:
        errors.append("Содержание – уберите абзацный отступ у заголовка")
    if abs(get_effective_left_indent(toc_header_para)) > 0.1:
        errors.append("Содержание – уберите отступ слева у заголовка")
    
    # Пустая строка после заголовка
    if toc_header_idx > 0 and toc_header_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[toc_header_idx + 1]
        if not is_empty_paragraph(next_para):
            errors.append("Содержание – добавьте пустую строку после заголовка")
    
# === 2. Сбор строк оглавления ===
    toc_lines = []
    is_inside_toc_zone = False
    toc_title_found = False

    # Шаг A: Пытаемся собрать строки внутри зоны автооглавления (по XML-тегам)
    for idx, p in enumerate(doc.paragraphs):
        txt = normalize_text(p.text)
        
        if not toc_title_found and txt.upper() in ["СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ"]:
            toc_title_found = True
            is_inside_toc_zone = True
            continue
            
        if is_inside_toc_zone:
            # Граница выхода из содержания
            if txt.upper() in ["ВВЕДЕНИЕ", "ВВЕДЕНIЕ"] or re.match(r'^(?:ГЛАВА|РАЗДЕЛ)\s+\d+', txt, re.IGNORECASE) or re.match(r'^1\s+[А-ЯЁ]', txt):
                is_inside_toc_zone = False
                break
                
            xml_text = "".join(node.text or "" for node in p._element.iter(qn('w:t'))).strip()
            p_text = xml_text if xml_text else p.text.strip()
            
            if p_text.strip():
                p.text = p_text  # Сохраняем текст в абзац для проверок
                toc_lines.append(p)
                
    # Шаг B: Если зона автооглавления не сработала (оглавление обычным текстом), используем бэкап-поиск по цифрам в конце
    if not toc_lines:
        start_search = toc_header_idx + 1 if (toc_header_idx is not None and toc_header_idx > 0) else 0
        for i in range(start_search, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            txt = p.text.strip()
            if not txt:
                continue
            # Стоп-слово: реальный заголовок в тексте без номера страницы
            if (txt.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"] or re.match(r'^\d+\.', txt)) and not re.search(r'\d+$', txt):
                break
            if re.search(r'\d+$', txt):
                toc_lines.append(p)

    toc_paragraphs = toc_lines 
        # === ЕСЛИ СТРОКИ НЕ НАЙДЕНЫ, ИЗВЛЕКАЕМ ИЗ ПОЛЯ TOC ===
    if not toc_lines:
        toc_lines_from_field = extract_toc_from_xml(doc)
        if toc_lines_from_field:
            toc_lines = toc_lines_from_field
        # Метод 2: Если строк с номерами страниц не нашли — ищем строки, похожие на пункты оглавления
    # (начинаются с цифр или являются спецразделами)
    if not toc_lines:
        for i in range(start_search, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            txt = p.text.strip()
            if not txt:
                continue
            
            # Стоп-слово
            if (txt.upper() in ["ВВЕДЕНИЕ", "ВВЕДЕНIЕ", "ЗАКЛЮЧЕНИЕ"] or 
                re.match(r'^(?:ГЛАВА|РАЗДЕЛ)\s+\d+', txt, re.IGNORECASE) or 
                re.match(r'^1\s+[А-ЯЁ]', txt)):
                break
            
            # Похоже на пункт оглавления: номер раздела или спецраздел
            if (re.match(r'^\d+(\.\d+)*\s+', txt) or 
                txt.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"]):
                toc_lines.append(p)        
        # Финальная попытка: прямое сканирование поля TOC во всём документе
    if not toc_lines:
        toc_field_paragraphs = []
        body_elems = list(doc.element.body)
        in_toc = False
        
        for elem in body_elems:
            if elem.tag == qn('w:p'):
                fldChar = elem.find('.//w:fldChar', NSMAP)
                if fldChar is not None and fldChar.get(qn('w:fldCharType')) == 'begin':
                    idx_in_body = body_elems.index(elem)
                    if idx_in_body + 1 < len(body_elems):
                        next_elem = body_elems[idx_in_body + 1]
                        instr = next_elem.find('.//w:instrText', NSMAP)
                        if instr is not None and instr.text and 'TOC' in instr.text.upper():
                            in_toc = True
                            continue
                
                if in_toc:
                    end_fldChar = elem.find('.//w:fldChar', NSMAP)
                    if end_fldChar is not None and end_fldChar.get(qn('w:fldCharType')) == 'end':
                        in_toc = False
                        continue
                    
                    for p in doc.paragraphs:
                        if p._element is elem:
                            txt = p.text.strip()
                            if txt and len(txt) > 3:
                                toc_field_paragraphs.append(p)
                            break
        
        if toc_field_paragraphs:
            toc_lines = toc_field_paragraphs
            st.success(f"✅ Найдено {len(toc_lines)} строк через финальное сканирование поля TOC")
       
    
        # Финальная попытка: прямое сканирование поля TOC во всём документе
    if not toc_lines:
        st.info("🔍 Запускаю финальное сканирование TOC по всему XML...")
        toc_field_paragraphs = []
        body_elems = list(doc.element.body)
        in_toc = False
        
        for elem in body_elems:
            if elem.tag == qn('w:p'):
                fldChar = elem.find('.//w:fldChar', NSMAP)
                if fldChar is not None and fldChar.get(qn('w:fldCharType')) == 'begin':
                    idx_in_body = body_elems.index(elem)
                    if idx_in_body + 1 < len(body_elems):
                        next_elem = body_elems[idx_in_body + 1]
                        instr = next_elem.find('.//w:instrText', NSMAP)
                        if instr is not None and instr.text and 'TOC' in instr.text.upper():
                            in_toc = True
                            st.text(f"✅ Найдено начало поля TOC в элементе {idx_in_body}")
                            continue
                
                if in_toc:
                    end_fldChar = elem.find('.//w:fldChar', NSMAP)
                    if end_fldChar is not None and end_fldChar.get(qn('w:fldCharType')) == 'end':
                        in_toc = False
                        st.text(f"🏁 Конец поля TOC в элементе {body_elems.index(elem)}")
                        continue
                    
                    for p in doc.paragraphs:
                        if p._element is elem:
                            txt = p.text.strip()
                            if txt and len(txt) > 3:
                                toc_field_paragraphs.append(p)
                                st.text(f"📄 Строка TOC: {repr(txt[:80])}")
                            break
        
        if toc_field_paragraphs:
            toc_lines = toc_field_paragraphs
            st.success(f"✅ Найдено {len(toc_lines)} строк через финальное сканирование поля TOC")
        else:
            st.error("❌ Финальное сканирование TOC не нашло строк. Поле TOC не обнаружено в документе.")
    
    # === ОТЛАДКА: выводим информацию о найденных строках ===
    st.markdown("---")
    st.markdown("### 🔍 Отладка сбора строк содержания")
    st.write(f"**toc_header_idx:** {toc_header_idx}")
    st.write(f"**start_search:** {start_search}")
    st.write(f"**Всего строк найдено в toc_lines:** {len(toc_lines)}")
    
    if toc_lines:
        with st.expander("Посмотреть найденные строки содержания"):
            for idx, p in enumerate(toc_lines):
                st.text(f"Строка {idx+1}: {repr(p.text[:100])}")
    else:
        st.warning("⚠️ toc_lines пуст — ни один метод не сработал")
        # Покажем, что находится между заголовком СОДЕРЖАНИЕ и ВВЕДЕНИЕ
        st.write("**Абзацы между СОДЕРЖАНИЕ и ВВЕДЕНИЕ:**")
        count = 0
        for i in range(start_search, min(start_search + 30, len(doc.paragraphs))):
            txt = doc.paragraphs[i].text
            if txt.strip():
                st.text(f"Абзац {i}: {repr(txt[:100])}")
                count += 1
                if count > 15:
                    st.text("... (показаны первые 15 непустых абзацев)")
                    break
    st.markdown("---")
       
    # === 3. Сбор заголовков из ТЕКСТА документа (для сверки) ===
    doc_headers = []
    in_bibliography = False
    
    for i, p in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
            
        txt = p.text.strip()
        if not txt or has_page_number(txt):
            continue
            
        if txt.upper() in ["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ЛИТЕРАТУРЫ", "БИБЛИОГРАФИЧЕСКИЙ СПИСОК"]:
            doc_headers.append(('special', txt.upper(), txt, False))
            in_bibliography = True
            continue
            
        if in_bibliography:
            continue
            
        if (re.match(r'^\d+\s+[А-ЯЁа-яё]', txt) or re.match(r'^\d+\.\s+[А-ЯЁа-яё]', txt)) and len(txt) < 100:
            num_match = re.match(r'^(\d+)[\.\s]', txt)
            if num_match:
                num = num_match.group(1)
                title = re.sub(r'^\d+[\.\s]*', '', txt).strip()
                doc_headers.append(('1', num, title, False))
        elif re.match(r'^\d+\.\d+\.?\s+[А-ЯЁа-яё]', txt) and len(txt) < 100:
            num_match = re.match(r'^(\d+\.\d+)(?:\.\d+)*', txt)
            if num_match:
                num = num_match.group(1)
                if len(num.split('.')) <= 2:
                    title = re.sub(r'^\d+\.\d+\s*', '', txt).strip()
                    doc_headers.append(('2', num, title, True))
        elif txt.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"]:
            doc_headers.append(('special', txt.upper(), txt, False))
    
           # === 4. Парсинг собранных строк оглавления ===
    toc_entries = []
    
    # Сначала пробуем улучшенный XML-метод
    toc_clean_strings = extract_toc_entries_clean(doc, toc_header_idx if toc_header_idx else 0)
    
    if toc_clean_strings:
        for entry_text in toc_clean_strings:
            # Ищем номер раздела в начале строки (например, "1.1 Методология обзора")
            # или "1 Логистика..."
            num_match = re.match(r'^(\d+(?:\.\d+)?)\s+', entry_text)
            if num_match:
                num = num_match.group(1)
                # Извлекаем название после номера
                title = re.sub(r'^\d+(?:\.\d+)?\s*', '', entry_text).strip()
                # Убираем возможные хвосты из цифр в конце названия
                title = re.sub(r'\s*\d+$', '', title)
                level = '1' if '.' not in num else '2'
                if title and len(title) > 1:
                    toc_entries.append((num, title, "?", level))
            else:
                # Специальные разделы без номера (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК ИСТОЧНИКОВ)
                # Убираем цифры страниц, которые могли прилипнуть
                clean_title = re.sub(r'\d+$', '', entry_text).strip()
                if clean_title.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"]:
                    toc_entries.append((None, clean_title, "?", 'special'))
                elif clean_title and len(clean_title) > 3:
                    # Возможно, это подраздел без номера или с нестандартным форматом
                    toc_entries.append((None, clean_title, "?", 'special'))
    
    # Резервный метод: если ничего не нашли, пробуем через toc_lines
    if not toc_entries and toc_lines:
        for p in toc_lines:
            txt = p.text.strip()
            if len(txt) > 150 or txt.startswith('•') or txt.startswith('-') or len(txt.split()) > 12:
                continue
            # Убираем номер страницы
            txt = re.sub(r'\s+\d+$', '', txt)
            txt = re.sub(r'\.{2,}', '', txt)
            
            page_match = re.search(r'(\d+)$', txt)
            if page_match:
                page_num = page_match.group(1)
                content = txt[:page_match.start()].strip()
            else:
                page_num = "?"
                content = txt
            
            num_match = re.match(r'^(\d+(?:\.\d+)?)[\.\s]', content)
            if num_match:
                num = num_match.group(1)
                title = re.sub(r'^\d+(?:\.\d+)?[\.\s]*', '', content).strip()
                title = re.sub(r'\s*\d+$', '', title)  # Убираем цифры в конце
                level = '1' if '.' not in num else '2'
                if len(num.split('.')) <= 2 and title:
                    toc_entries.append((num, title, page_num, level))
            elif content.strip():
                clean_content = re.sub(r'\d+$', '', content).strip()
                if clean_content:
                    toc_entries.append((None, clean_content, page_num, 'special'))
    
    # Удаляем дубликаты
    seen = set()
    unique_toc_entries = []
    for entry in toc_entries:
        key = entry[0] if entry[0] else entry[1]
        if key not in seen:
            seen.add(key)
            unique_toc_entries.append(entry)
    toc_entries = unique_toc_entries
    
    # Отладка: выводим обработанные записи
    st.write(f"**Обработано записей оглавления:** {len(toc_entries)}")
    for entry in toc_entries:
        st.text(f"  {entry}")
    
    # === 5. Сверка Содержания и Текста ===
    doc_headers_by_num = {}
    for level, num, title, is_sub in doc_headers:
        if level in ('1', '2'):
            doc_headers_by_num[num] = title
        elif level == 'special':
            doc_headers_by_num[title] = title  
            
    normalized_doc_headers = {}
    for key, value in doc_headers_by_num.items():
        norm_key = key.rstrip('.')
        normalized_doc_headers[norm_key] = value
    doc_headers_by_num = normalized_doc_headers
    
    # Сверка 1: Все ли пункты из Оглавления присутствуют в самом тексте
    for num, title, page, level in toc_entries:
        key = num if num else title
        key = key.rstrip('.')
        if key not in doc_headers_by_num:
            continue
        expected_title = doc_headers_by_num[key]
        if title.lower() != expected_title.lower():
            if re.sub(r'^(\d+\.?)\s*', '', title).lower() != re.sub(r'^(\d+\.?)\s*', '', expected_title).lower():
                errors.append(f"Содержание – заголовок «{title}» не соответствует заголовку в тексте («{expected_title}»)")
    
    # Сверка 2: Все ли реальные разделы текста внесены в Оглавление
    for level, num, title, is_sub in doc_headers:
        key = num if level in ('1', '2') else title
        key = key.rstrip('.')
        
        found = False
        for entry in toc_entries:
            entry_key = entry[0] if entry[0] else entry[1]
            if entry_key:
                entry_key = entry_key.rstrip('.')
                if entry_key.lower() == key.lower():
                    found = True
                    break
                    
        if not found:
            errors.append(f"Содержание – отсутствует раздел «{title}» (хотя он есть в тексте документа)")

    # === 6. Проверка форматирования строк оглавления ===
    for p in toc_lines:
        txt = p.text.strip()
        if not txt: continue
        
        # Пропускаем проверку стилей для отфильтрованного мусора
        if len(txt) > 150 or txt.startswith('•') or len(txt.split()) > 12:
            continue
            
        sizes = get_font_size_pt(p)
        if sizes and any(abs(s - 14) > 0.5 for s in sizes):
            errors.append("Содержание – установите размер шрифта 14 пт для всех строк")
            break
        if is_paragraph_bold(p):
            errors.append("Содержание – уберите полужирное начертание")
            break
        if get_effective_alignment(p) != WD_ALIGN_PARAGRAPH.JUSTIFY:
            errors.append("Содержание – выровняйте строки по ширине")
            break
        if abs(get_effective_first_line_indent(p)) > 0.1:
            errors.append("Содержание – уберите абзацный отступ")
            break
        if abs(get_effective_left_indent(p)) > 0.1:
            errors.append("Содержание – уберите отступ слева")
            break
            
        try:
            line_spacing = p.paragraph_format.line_spacing
            if line_spacing and abs(line_spacing - 1.2) > 0.05:
                errors.append("Содержание – установите междустрочный интервал множитель 1,2")
                break
        except:
            pass

    # ==============================================================================
    # ТЕСТОВЫЙ БЛОК ОТЛАДКИ С ТЕКУЩИМИ ДАННЫМИ
    # ==============================================================================
    st.markdown("---")
    st.markdown("### 🔍 Отладочная информация оглавления")
    if 'toc_lines' in locals() and toc_lines:
        st.write(f"**Всего строк найдено в зоне TOC:** {len(toc_lines)}")
        with st.expander("Посмотреть строки Содержания"):
            for idx, p in enumerate(toc_lines):
                st.text(f"Строка {idx+1}: {repr(p.text)}")
    
    if 'toc_entries' in locals() and toc_entries:
        st.write(f"**Успешно распознано пунктов (toc_entries):** {len(toc_entries)}")
        with st.expander("Посмотреть распознанные пункты оглавления"):
            for entry in toc_entries:
                st.code(str(entry))
    else:
        st.warning("⚠️ Список toc_entries пуст.")
        
    if 'doc_headers' in locals() and doc_headers:
        st.write(f"**Найденные заголовки в теле документа:** {len(doc_headers)}")
        with st.expander("Посмотреть заголовки из тела"):
            for idx, h_item in enumerate(doc_headers):
                st.text(f"Заголовок {idx+1}: {repr(h_item)}")
    st.markdown("---")
    # ==============================================================================
   
        # === УДАЛЕНИЕ ЛОЖНЫХ ОШИБОК ОБ ОТСУТСТВИИ РАЗДЕЛОВ ===
    errors = [err for err in errors if not (err.startswith("Содержание – отсутствует раздел") and "хотя он есть в тексте документа" in err)]
    return errors


# ------------------------------------------------------------
# Главная проверка документа
# ------------------------------------------------------------
def check_word_document(file):
    doc = docx.Document(file)
    auto_issues = []
    manual_checks = []
    manual_issues = manual_checks

    # --- ПОЛЯ ---
    margins_ok = True
    for section in doc.sections:
        if (abs(section.left_margin.mm - 20) > 0.5 or
            abs(section.right_margin.mm - 20) > 0.5 or
            abs(section.top_margin.mm - 20) > 0.5 or
            abs(section.bottom_margin.mm - 20) > 0.5):
            margins_ok = False
            break
    if not margins_ok:
        auto_issues.append("Поля страниц – установите левое 20 мм, правое 20 мм, верхнее 20 мм, нижнее 20 мм")

    # --- ПОИСК НАЧАЛА ОСНОВНОГО ТЕКСТА ---
    start_idx = None
    intro_found = False
   # auto_issues = []

    # Ищем "ВВЕДЕНИЕ"
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        if has_page_number(txt):
            continue
        if txt.upper() == "ВВЕДЕНИЕ" and is_section_header(txt):
            intro_found = True
            start_idx = i
            break

    # Если не нашли "ВВЕДЕНИЕ" — ищем любой раздел
    if not intro_found:
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            if has_page_number(txt):
                continue
            if is_section_header(txt):
                start_idx = i
                break
        
        if start_idx is not None:
            auto_issues.append("Отсутствует введение, оформленное как заголовок. Проверка начинается с первого найденного раздела.")

    # Если нет ни одного раздела
    if start_idx is None:
        return ["Отсутствует введение. Документ не содержит заголовков разделов."]

    # --- ПРОВЕРКА СОДЕРЖАНИЯ (должна быть первой) ---
    toc_errors = check_toc(doc, start_idx)
   # st.write(f"DEBUG: toc_errors = {toc_errors}")  # временно
    # Добавляем ошибки содержания в начало списка auto_issues
    auto_issues = toc_errors + auto_issues
    #st.write(f"DEBUG: auto_issues после содержания = {auto_issues}")  # временно

    
    # --- НУМЕРАЦИЯ СТРАНИЦ ---
    try:
        intro_body_idx = None
        body_elements = list(doc.element.body)
        for i, elem in enumerate(body_elements):
            if elem.tag == qn('w:p'):
                full_text = ''.join(t.text or '' for t in elem.findall('.//w:t', qn('w'))).strip().upper()
                if full_text == 'ВВЕДЕНИЕ':
                    intro_body_idx = i
                    break
        if intro_body_idx is not None:
            page_issues = check_page_numbering(file, intro_body_idx)
            auto_issues.extend(page_issues)
    except:
        pass

    toc_entries = extract_toc_entries(doc, start_idx)

    # --- ГРАНИЦЫ СПИСКА ИСТОЧНИКОВ ---
    lit_start = None
    for i in range(start_idx, len(doc.paragraphs)):
        txt = doc.paragraphs[i].text.strip()
        txt_upper = txt.upper()
        if txt_upper in ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ") and not has_page_number(txt):
            lit_start = i
            if txt_upper == "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ":
                auto_issues.append("Список источников – замените заголовок на «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ»")
            break   
    
    lit_end = len(doc.paragraphs)
    if lit_start is not None:
        appendix_keywords = ["ПРИЛОЖЕНИЕ", "ПРИЛОЖЕНИЯ", "APPENDIX"]
        for i in range(lit_start + 1, len(doc.paragraphs)):
            txt = doc.paragraphs[i].text.strip().upper()
            if any(txt.startswith(kw) for kw in appendix_keywords):
                lit_end = i
                break

    # --- ПРОВЕРКА ОСНОВНОГО ТЕКСТА ---
    figure_counter = 0
    prev_para_empty = False
    prev_was_formula = False
    prev_was_section_header = False
    prev_was_subsection = False
    end_idx = lit_start if lit_start is not None else len(doc.paragraphs)
    list_errors = []
    indent_issues = []

    figure_numbers_found = []
    table_numbers_found = []

    para_to_body_idx = {}
    body_elems = list(doc.element.body)
    for i, elem in enumerate(body_elems):
        if elem.tag == qn('w:p'):
            for j, p in enumerate(doc.paragraphs):
                if p._element is elem:
                    para_to_body_idx[j] = i
                    break

    table_captions_info = []

    try:
        start_body_pos = para_to_body_idx[start_idx]
    except:
        start_body_pos = 0

    for idx in range(start_idx, end_idx):
        p = doc.paragraphs[idx]
        # Читаем текст напрямую из XML, чтобы не потерять заголовки-ссылки
        text = "".join(node.text or "" for node in p._element.iter(qn('w:t'))).strip()
        norm_text = normalize_text(text)
        if len(norm_text) < 5 or re.match(r'^[,.\(\)\d\s]+$', norm_text):
            continue
        style_name = p.style.name.lower() if p.style else ""

        if not text:
            prev_para_empty = True
            if list_errors:
                first_text = list_errors[0][1]
                auto_issues.append(f"Список начиная с «{first_text[:50]}» и далее – замените круглый маркер (•) на тире, букву или цифру")
                list_errors = []
            continue

        if has_page_number(text):
            prev_para_empty = False
            continue

        # --- ОПРЕДЕЛЕНИЕ ЗАГОЛОВКОВ 1 УРОВНЯ ---
        is_level1 = False
        if 'heading 1' in style_name or 'заголовок 1' in style_name:
            is_level1 = True
        elif text.upper() in ["ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"]:
            is_level1 = True
        elif re.match(r'^[1-9]\d*\s+[А-ЯA-Z]', text):
            is_level1 = True
        elif is_section_header(norm_text):
            is_level1 = True

        # --- ОПРЕДЕЛЕНИЕ ПОДЗАГОЛОВКОВ (1.1, 1.2, 1.1.1) ---
        is_subsection = False
        if 'heading 2' in style_name or 'заголовок 2' in style_name or 'heading 3' in style_name or 'заголовок 3' in style_name:
            is_subsection = True
        elif re.match(r'^\d+(\.\d+)+\.?\s+', text):
            is_subsection = True

        # Корректировка: подраздел не может быть разделом 1 уровня
        if is_subsection:
            is_level1 = False
        
        # Флаги для предудущих состояний
        if is_level1 or is_subsection:
            if is_level1:
                prev_was_section_header = True
                prev_was_subsection = False
            else:
                prev_was_subsection = True
                prev_was_section_header = False
        else:
            prev_was_section_header = False
            prev_was_subsection = False

        # Для подзаголовков: проверяем, что после нет пустой строки
        if is_subsection and idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[idx + 1]
            if is_empty_paragraph(next_p):
                auto_issues.append(f"Подраздел «{text[:50]}» – после подзаголовка не должно быть пустой строки")

        # Флаги для предудущих состояний
        if is_level1 or is_subsection:
            if is_level1:
                prev_was_section_header = True
                prev_was_subsection = False
            else:
                prev_was_subsection = True
                prev_was_section_header = False
        else:
            prev_was_section_header = False
            prev_was_subsection = False

        # --- ЛОГИКА ДЛЯ НЕ-ЗАГОЛОВКОВ (списки, формулы, рисунки, основной текст) ---
        if not is_level1 and not is_subsection:
            # Проверка списков
            is_list, marker_type, marker_valid = get_list_marker_info(p, doc)
            if is_list:
                if not marker_valid:
                    if list_errors and idx == list_errors[-1][0] + 1:
                        list_errors.append((idx, text, marker_type))
                    else:
                        if list_errors:
                            first_text = list_errors[0][1]
                            auto_issues.append(f"Список начиная с «{first_text[:50]}» и далее – замените круглый маркер (•) на тире, букву или цифру")
                            list_errors = []
                        list_errors.append((idx, text, marker_type))
                else:
                    if list_errors:
                        first_text = list_errors[0][1]
                        auto_issues.append(f"Список начиная с «{first_text[:50]}» и далее – замените круглый маркер (•) на тире, букву или цифру")
                        list_errors = []
                prev_para_empty = False
                continue

            if list_errors:
                first_text = list_errors[0][1]
                auto_issues.append(f"Список начиная с «{first_text[:50]}» и далее – замените круглый маркер (•) на тире, букву или цифру")
                list_errors = []

            # Продолжение таблицы
            if is_table_continuation(norm_text):
                first_line = get_effective_first_line_indent(p)
                # Если отступ не задан локально (унаследован из стиля), считаем его нулевым
                if p.paragraph_format.first_line_indent is None:
                    first_line = 0.0
                if abs(first_line) > 0.3:
                    auto_issues.append(f"«{text[:50]}» – уберите абзацный отступ (должен быть 0 см)")
                prev_para_empty = False
                continue

            # Пояснение к формуле
            if is_formula_where_line(norm_text):
                errors = check_formula_explanation(text, p, prev_was_formula, prev_para_empty)
                auto_issues.extend(errors)
                prev_para_empty = False
                continue

            # Формула
            if is_formula_or_equation(norm_text):
                prev_was_formula = True
                prev_para_empty = False
                continue

            # Проверка рисунков
            is_figure = norm_text.startswith("Рисунок") or norm_text.startswith("Рис.")
            is_table_caption = norm_text.startswith("Таблица")

            if is_figure:
                num_match = re.search(r'(?:Рисунок|Рис\.)\s*:?\s*(\d+(?:\.\d+)?)', norm_text)
                if num_match:
                    fig_num = num_match.group(1)
                    figure_numbers_found.append(float(fig_num))
                else:
                    figure_counter += 1
                    fig_num = str(figure_counter)
                fig_number = f"Рисунок {fig_num}"

                if norm_text.startswith("Рис."):
                    auto_issues.append(f"{fig_number} – измените «Рис.» на «Рисунок»")
                if get_effective_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    auto_issues.append(f"{fig_number} – выровняйте подпись по центру")
                if text.endswith(".") and not re.search(r'\([^)]*\)\.$', text):
                    auto_issues.append(f"{fig_number} – удалите точку в конце")
                m = re.match(r'^(?:Рисунок|Рис\.)\s+\d+(?:\.\d+)?\s*[–—]\s*(.+)$', norm_text)
                if m:
                    title = m.group(1).strip()
                    if title and title[0].islower():
                        auto_issues.append(f"{fig_number} – название должно начинаться с большой буквы")

                sizes = get_font_size_pt(p)
                if sizes:
                    if any(abs(s - 14) > 0.5 for s in sizes):
                        auto_issues.append(f"{fig_number} – установите размер шрифта 14 пт (сейчас {', '.join(str(s) for s in sizes)} пт)")

                body_idx = para_to_body_idx.get(idx)
                if body_idx is not None:
                    empty_errors = check_empty_line_before_after(doc, body_idx, start_body_pos, fig_number)
                    auto_issues.extend(empty_errors)
                prev_para_empty = False
                first_line = get_effective_first_line_indent(p)
                if abs(first_line) > 0.1:
                    auto_issues.append(f"{fig_number} – уберите абзацный отступ (должен быть 0 см, сейчас {first_line:.1f} см)")

                
                continue

            # Подпись таблицы
            elif is_table_caption and not is_table_continuation(norm_text):
                tbl_match = re.match(r'Таблица\s+(\d+(?:\.\d+)?)', norm_text)
                if tbl_match:
                    tbl_num = tbl_match.group(1)
                    key = f"Таблица {tbl_num}"

                    if not re.match(r'Таблица\s+\d+\s+[-–—]{1,2}\s+', text):
                        auto_issues.append(f"Таблица {tbl_num} – используйте тире между номером и названием (например, «Таблица 5 – Название»)")

                    if text.rstrip().endswith("."):
                        auto_issues.append(f"{key} – удалите точку в конце названия")
                    sizes = get_font_size_pt(p)
                    if sizes:
                        if any(abs(s - 14) > 0.5 for s in sizes):
                            auto_issues.append(f"{key} – установите размер шрифта 14 пт (сейчас {', '.join(str(s) for s in sizes)} пт)")

                prev_para_empty = False
                continue

            # Основной текст
            key = norm_text[:50]
            first_line = get_effective_first_line_indent(p)
            if abs(first_line - 1.0) > 0.2:
                indent_issues.append((key, first_line))
            if p.paragraph_format.space_before and p.paragraph_format.space_before.pt > 0.5:
                auto_issues.append(f"«{key}» – интервал перед абзацем должен быть 0 пт")

        # === Проверка заголовков 1 уровня ===
        if is_level1:
            key = text[:80]
            if text.upper() != "ВВЕДЕНИЕ":
                body_idx = para_to_body_idx.get(idx)
                if body_idx is not None and not is_on_new_page(doc, body_idx, start_body_pos):
                    auto_issues.append(f"«{key}» – раздел должен начинаться с новой страницы")
            first_line = get_effective_first_line_indent(p)
            if p.paragraph_format.first_line_indent is None:
                first_line = 0.0
            if abs(first_line) > 0.3:
                auto_issues.append(f"«{key}» – уберите абзацный отступ у заголовка")
            if not is_paragraph_bold(p):
                auto_issues.append(f"«{key}» – заголовок раздела должен быть полужирным")
            if re.match(r'^\d+\.', norm_text) and not is_all_caps(norm_text):
                auto_issues.append(f"«{key}» – заголовок раздела должен быть прописными буквами")
            if get_effective_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                auto_issues.append(f"«{key}» – выровняйте заголовок по центру")
            if text.endswith("."):
                auto_issues.append(f"«{key}» – удалите точку в конце")

            if not has_proper_spacing_after_header(doc, idx):
                auto_issues.append(f"«{key}» – после заголовка должна быть пустая строка (или интервал после ≥ 12 pt)")

            if text.upper() == "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ":
                auto_issues.append(f"«{text[:50]}» – замените на «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ»")

        # === Проверка подзаголовков ===
        elif is_subsection:
            sub_name = re.sub(r'^\d+\.\d+(\.\d+)?\s*', '', norm_text).strip()
            key = f"Подраздел «{sub_name[:50]}»"
            first_line = get_effective_first_line_indent(p)

            # Если отступ не задан локально (унаследован из стиля), пропускаем проверку
            if p.paragraph_format.first_line_indent is None:
                first_line = 1.0   # принудительно устанавливаем правильное значение
            
            if abs(first_line - 1.0) > 0.2:
                auto_issues.append(f"{key} – установите абзацный отступ 1,0 см (сейчас {first_line:.1f} см)")
            if not is_paragraph_bold(p):
                auto_issues.append(f"{key} – заголовок должен быть полужирным")
            if get_effective_alignment(p) != WD_ALIGN_PARAGRAPH.JUSTIFY:
                auto_issues.append(f"{key} – выровняйте по ширине")
            if text.endswith("."):
                auto_issues.append(f"{key} – удалите точку в конце")

            # Проверка пустой строки перед подразделом
            if idx > 0:
                prev_p = doc.paragraphs[idx - 1]
                prev_txt = prev_p.text.strip()
                if prev_txt != "":
                    try:
                        body_elems_local = list(doc.element.body)
                        body_idx = body_elems_local.index(p._element)
                    except:
                        body_idx = -1
                    if body_idx != -1 and is_on_new_page(doc, body_idx, start_body_pos=start_idx):
                        pass
                    elif is_section_header(prev_txt) or re.match(r'^\d+\.\d+', prev_txt):
                        auto_issues.append(f"{key} – добавьте пустую строку перед подразделом")
                    else:
                        manual_issues.append(
                            f"{key} – проверьте визуально: если он естественно перенесся на новую страницу, то отступ перед ним не нужен. "
                            f"Если он идет внутри страницы, добавьте пустую строку."
                        )

        # Сброс флагов для следующей итерации
        if text:
            prev_para_empty = False
            prev_was_formula = False
        else:
            prev_para_empty = True

    # --- ЗАВЕРШАЮЩИЕ ПРОВЕРКИ ---
    if indent_issues:
        if len(indent_issues) > 2:
            first_key, first_indent = indent_issues[0]
            auto_issues.append(
                f"Основной текст, начиная со строки «{first_key}» – "
                f"установите абзацный отступ 1,0 см (сейчас {first_indent:.1f} см)"
            )
        else:
            for key, fl in indent_issues:
                auto_issues.append(f"«{key}» – установите абзацный отступ 1,0 см (сейчас {fl:.1f} см)")

    if list_errors:
        first_text = list_errors[0][1]
        auto_issues.append(f"Список начиная с «{first_text[:50]}» и далее – замените круглый маркер (•) на тире, букву или цифру")

    if figure_numbers_found:
        int_nums = sorted(set(int(n) for n in figure_numbers_found if n == int(n)))
        if int_nums:
            if int_nums[0] != 1:
                auto_issues.append("Рисунки – нумерация должна начинаться с 1")
            else:
                for expected in range(1, int_nums[-1] + 1):
                    if expected not in int_nums:
                        auto_issues.append(f"Рисунки – пропущен рисунок {expected}")
                        break

    table_seq_issues = []
    if table_numbers_found:
        int_tbl_nums = sorted(set(int(n) for n in table_numbers_found if n == int(n)))
        if int_tbl_nums:
            if int_tbl_nums[0] != 1 or any(expected not in int_tbl_nums for expected in range(1, int_tbl_nums[-1] + 1)):
                table_seq_issues.append("Таблицы – неверная нумерация")

    # --- ТАБЛИЦЫ ---
       # --- ТАБЛИЦЫ ---
        # --- НОВАЯ ПРОВЕРКА ТАБЛИЦ (расположение названия и пустая строка после) ---
    # Определяем регулярное выражение для названия таблицы
    table_title_re = re.compile(r'^таблица\s+\d+', re.IGNORECASE)
    TABLE_CONTINUATION_RE = re.compile(r'^(продолжение|окончание)\s+таблицы', re.IGNORECASE)

    # Получаем плоский список элементов (параграфы и таблицы) в порядке следования
    try:
        elements = list(iter_block_items(doc))
    except Exception as e:
        # Если функция iter_block_items не определена (защита от ошибки)
        manual_checks.append(f"Ошибка при обходе элементов документа: {e}")
        elements = []

    for idx, item in enumerate(elements):
        # Если текущий элемент — ТАБЛИЦА
        if isinstance(item, docx.table.Table):
            # --- 1. Проверяем расположение названия (должно быть ПЕРЕД таблицей) ---
            has_title_before = False
            prev_text = ""
            if idx > 0 and isinstance(elements[idx - 1], docx.text.paragraph.Paragraph):
                prev_text = normalize_text(elements[idx - 1].text)
                if table_title_re.match(prev_text):
                    has_title_before = True

            has_title_after = False
            next_text = ""
            if idx + 1 < len(elements) and isinstance(elements[idx + 1], docx.text.paragraph.Paragraph):
                next_text = normalize_text(elements[idx + 1].text)
                if table_title_re.match(next_text):
                    has_title_after = True

            # Извлекаем номер таблицы для сообщений
            table_num_match = re.search(r'\d+', prev_text if has_title_before else (next_text if has_title_after else ""))
            table_num = table_num_match.group(0) if table_num_match else "?"

            if has_title_after and not has_title_before:
                auto_issues.append(f"Таблица {table_num} – название должно быть ПЕРЕД таблицей (сейчас расположено после)")
            elif not has_title_before and not has_title_after:
                # Название не найдено – помечаем для ручной проверки
                manual_checks.append(f"📋 Таблица (№ {table_num} по счёту) – не обнаружено стандартное название формата 'Таблица X'. Проверьте вручную.")

            # --- 2. Проверяем, есть ли пустая строка ПОСЛЕ таблицы (перед следующим текстом) ---
            # Сдвигаем индекс, чтобы пропустить возможное название, если оно ошибочно идёт после таблицы
            check_idx = idx + 2 if has_title_after else idx + 1

            
            if check_idx < len(elements):
                next_item = elements[check_idx]
                if isinstance(next_item, docx.text.paragraph.Paragraph):
                    raw_text = next_item.text.strip()
        # Проверяем: если текст пуст ИЛИ это просто набор пробелов/невидимок (len < 2)
        # Мы считаем такой абзац "пустой строкой" (правильно)
                    is_actually_empty = len(raw_text) < 2
        
        # Если абзац НЕ пустой И НЕ является "Продолжение таблицы"
                    if not is_actually_empty and not TABLE_CONTINUATION_RE.match(raw_text):
                        short_next = (raw_text[:40] + "...") if len(raw_text) > 40 else raw_text
                        manual_checks.append(
                            f"📋 После таблицы {table_num} отсутствует пустая строка перед текстом «{short_next}»."
            )

    # Проверка нумерации таблиц (оставляем старую логику, если нужна)
    if table_numbers_found:
        int_tbl_nums = sorted(set(int(n) for n in table_numbers_found if n == int(n)))
        if int_tbl_nums:
            if int_tbl_nums[0] != 1 or any(expected not in int_tbl_nums for expected in range(1, int_tbl_nums[-1] + 1)):
                auto_issues.append("Таблицы – неверная нумерация")

    # --- СПИСОК ИСТОЧНИКОВ ---
    if lit_start is not None:
        sources_with_issues = 0
        for i in range(lit_start + 1, lit_end):
            source = doc.paragraphs[i]
            txt = source.text.strip()
            if not txt or has_page_number(txt):
                continue
            has_issue = False
            left_indent = get_effective_left_indent(source)
            if abs(left_indent) > 0.1:
                has_issue = True
            first_line = get_effective_first_line_indent(source)
            if abs(first_line - 1.0) > 0.1:
                has_issue = True
            if has_issue:
                sources_with_issues += 1
        if sources_with_issues > 0:
            auto_issues.append(
                "Список источников – проверьте оформление: "
                "отступ слева 0 см, отступ первой строки 1,0 см, "
                "междустрочный интервал 1,2 (множитель), выравнивание по ширине"
            )

    # --- ОБЪЕДИНЯЕМ ВСЕ РУЧНЫЕ ПРОВЕРКИ ---
    # manual_issues (из подзаголовков) и manual_checks (из таблиц/рисунков)
    all_manual = []
    if manual_issues:
        all_manual.extend(manual_issues)
    if manual_checks:
        all_manual.extend(manual_checks)
    
    # Добавляем глобальные напоминания
    all_manual.append("📌 ПРОВЕРЬТЕ ВРУЧНУЮ: титульный лист, задание, содержание, введение (первые 4 страницы) на соответствие шаблону.")
    all_manual.append("⚠️ ВНИМАНИЕ: Проверьте вручную наличие надписей «Продолжение таблицы» или «Окончание таблицы» на каждой странице, где таблица разрывается.")
    all_manual.append("⚠️ ВНИМАНИЕ: Проверьте правильность переноса рисунков на новую страницу и их подписей.")
    
    # Формируем итоговый список
    all_issues = auto_issues
    if all_manual:
        all_issues.append("📋 Для проверки человеком:")
        all_issues.extend(all_manual)
    
    return group_issues(all_issues)

# --- ИНТЕРФЕЙС STREAMLIT ---
st.set_page_config(page_title="Нормоконтроль документов", layout="centered")
st.title("📊 Автоматическая проверка документов Word")
st.write("Загрузите документ в формате .docx – проверка по полному чек-листу.")
uploaded_file = st.file_uploader("Выберите файл", type=["docx"])

if uploaded_file is not None:
    with st.spinner("Проверяем..."):
        results = check_word_document(uploaded_file)
    
    st.subheader("Результаты проверки:")
    
    # Разделяем автоматические ошибки и ручные проверки
    manual_mode = False
    auto_errors = []
    manual_items = []
    
    for r in results:
        if r.startswith("📋 Для проверки человеком:"):
            manual_mode = True
            continue
        if manual_mode:
            manual_items.append(r)
        else:
            auto_errors.append(r)
    
    # Выводим автоматические ошибки
    if auto_errors:
        for err in auto_errors:
            st.write(f"• {err}")
    else:
        st.success("✅ Автоматических ошибок не найдено.")
    
    # ВСЕГДА выводим блок для человека (3 обязательных пункта)
    st.markdown("---")
    #st.markdown("### 📋 Экспертная проверка (ОБЯЗАТЕЛЬНО):")
    st.markdown("**📋 Экспертная проверка (ОБЯЗАТЕЛЬНО)**") 
    st.markdown("Проверьте титульный лист, задание, содержание, введение (первые 4 страницы) на соответствие шаблону.")
    st.markdown("Проверьте вручную наличие надписей «Продолжение таблицы» / «Окончание таблицы» на каждой странице разрыва.")
    st.markdown("Проверьте правильность переноса рисунков на новую страницу и их подписей.")
    st.markdown("Проверьте, что оглавление является автособираемым")
        
    # Если есть дополнительные замечания из manual_items (от кода), выводим их
    if manual_items:
        st.markdown("#### Дополнительные замечания:")
        for item in manual_items:
            if item.strip():
                st.write(f"• {item}")
    
    # Финальное сообщение
    st.info("💬 Сообщение для человека: проверьте, пожалуйста, машине не всегда можно доверять")
